"""
DevOps Runbook Assistant — FastAPI Backend
RAG-Powered Intelligent System for Automated Incident Response
"""
from fastapi import FastAPI, HTTPException, UploadFile, File, Request
from fastapi.staticfiles import StaticFiles
from fastapi.responses import HTMLResponse, RedirectResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import Optional
import uvicorn, io, os

from database import (
    init_db, get_chat_history, save_chat,
    register_user, get_all_users, set_user_active,
    update_user_role, delete_user, count_users
)
init_db()

from rag import RAGPipeline
from auth import authenticate_user, create_session, get_current_user
from audit import log_query, get_logs, get_metrics, submit_feedback

app = FastAPI(title="DevOps Runbook Assistant", version="4.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"], allow_credentials=True,
    allow_methods=["*"], allow_headers=["*"],
)

rag = RAGPipeline()

FRONTEND_BUILD = os.path.join(os.path.dirname(__file__), "frontend", "build")
if os.path.exists(FRONTEND_BUILD):
    assets = os.path.join(FRONTEND_BUILD, "assets")
    if os.path.exists(assets):
        app.mount("/assets", StaticFiles(directory=assets), name="assets")


# ── Pydantic Models ────────────────────────────────────────────────────────
class LoginReq(BaseModel):
    username: str
    password: str

class RegisterReq(BaseModel):
    username: str
    password: str
    role: Optional[str] = "devops"
    session_token: Optional[str] = None

class QueryReq(BaseModel):
    question: str
    session_token: str
    file_content: Optional[str] = None
    file_name: Optional[str] = None

class FeedbackReq(BaseModel):
    log_id: int
    feedback: str
    session_token: str

class IngestReq(BaseModel):
    title: str
    content: str
    tags: Optional[str] = ""

class UserRoleReq(BaseModel):
    session_token: str
    role: str

class UserActiveReq(BaseModel):
    session_token: str
    is_active: bool


# ── SPA Serve ─────────────────────────────────────────────────────────────
def _spa():
    index = os.path.join(FRONTEND_BUILD, "index.html")
    if os.path.exists(index):
        with open(index) as f:
            return HTMLResponse(f.read())
    return HTMLResponse("""<!DOCTYPE html><html><body style="font-family:sans-serif;padding:40px;background:#0f172a;color:#e2e8f0">
    <h2>⚙ DevOps Runbook Assistant — Build frontend first</h2>
    <pre style="background:#1e293b;padding:16px;border-radius:8px;color:#93c5fd">cd frontend && npm install && npm run build</pre>
    <p>Then restart the server.</p></body></html>""")

@app.get("/")
async def root():
    return RedirectResponse("/login")


@app.get("/login")
async def login_page():
    return _spa()


@app.get("/chat")
async def chat_page():
    return _spa()


@app.get("/logs")
async def logs_page():
    return _spa()


@app.get("/admin")
async def admin_page():
    return _spa()


@app.get("/register")
async def register_page():
    return _spa()


# ── Auth ───────────────────────────────────────────────────────────────────
@app.post("/api/login")
async def login(req: LoginReq):
    user = authenticate_user(req.username, req.password)
    if not user: raise HTTPException(401, "Invalid username or password")
    return {"token": create_session(user["id"]), "username": user["username"], "role": user["role"]}

# ── Auth ───────────────────────────────────────────────────────────────────
@app.post("/api/login")
async def login(req: LoginReq):
    user = authenticate_user(req.username, req.password)

    if not user:
        raise HTTPException(
            status_code=401,
            detail="Invalid username or password"
        )

    return {
        "token": create_session(user["id"]),
        "username": user["username"],
        "role": user["role"]
    }


@app.post("/api/register")
async def register(req: RegisterReq):
    try:
        role = req.role if req.role else "admin"

        register_user(
            req.username,
            req.password,
            role
        )

        return {
            "success": True,
            "message": f"✅ User '{req.username}' registered successfully as {role}"
        }

    except ValueError as e:
        raise HTTPException(
            status_code=400,
            detail=str(e)
        )

    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=str(e)
        )
        

# ── File Upload ────────────────────────────────────────────────────────────
@app.post("/api/upload")
async def upload(file: UploadFile = File(...)):
    content  = await file.read()
    filename = file.filename.lower()
    text     = ""
    try:
        if filename.endswith(".pdf"):
            try:
                import PyPDF2
                reader = PyPDF2.PdfReader(io.BytesIO(content))
                for page in reader.pages:
                    pt = page.extract_text()
                    if pt: text += pt + "\n"
            except Exception as e:
                print(f"PyPDF2 error: {e}")
            if not text.strip():
                try:
                    from pdf2image import convert_from_bytes
                    import pytesseract
                    images = convert_from_bytes(content, dpi=200)
                    for img in images:
                        text += pytesseract.image_to_string(img, lang="eng") + "\n"
                except Exception:
                    text = "PDF_SCANNED: Convert to .docx for best results."
        elif filename.endswith((".docx", ".doc")):
            try:
                import docx
                doc = docx.Document(io.BytesIO(content))
                parts = [p.text for p in doc.paragraphs if p.text.strip()]
                for table in doc.tables:
                    for row in table.rows:
                        rt = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                        if rt: parts.append(rt)
                text = "\n".join(parts)
            except Exception:
                text = content.decode("utf-8", errors="ignore")
        elif filename.endswith((".xlsx", ".xls")):
            try:
                import pandas as pd
                xl = pd.ExcelFile(io.BytesIO(content))
                sheets = []
                for sheet in xl.sheet_names:
                    df = pd.read_excel(io.BytesIO(content), sheet_name=sheet)
                    sheets.append(f"=== Sheet: {sheet} ===\n{df.to_string()}")
                text = "\n\n".join(sheets)
            except Exception as e:
                text = f"Excel error: {e}"
        elif filename.endswith(".csv"):
            try:
                import pandas as pd
                df = pd.read_csv(io.BytesIO(content))
                text = df.to_string()
            except Exception:
                text = content.decode("utf-8", errors="ignore")
        elif filename.endswith(".json"):
            try:
                import json as jsonlib
                data = jsonlib.loads(content.decode("utf-8"))
                text = jsonlib.dumps(data, indent=2)
            except Exception:
                text = content.decode("utf-8", errors="ignore")
        else:
            text = content.decode("utf-8", errors="ignore")
    except Exception as e:
        text = content.decode("utf-8", errors="ignore")

    text = text.strip() or "File empty or unreadable."
    return {"success": True, "filename": file.filename, "content": text[:8000], "size": len(content), "chars": len(text)}


# ── Query ──────────────────────────────────────────────────────────────────
@app.post("/api/query")
async def query(req: QueryReq):
    user = get_current_user(req.session_token)
    if not user: raise HTTPException(401, "Invalid or expired session")

    result = await rag.query(
        question     = req.question,
        file_content = req.file_content,
        file_name    = req.file_name,
    )

    save_chat(user["id"], req.session_token, "user",      req.question)
    save_chat(user["id"], req.session_token, "assistant", result.get("answer", ""))

    log_id = log_query(
        user["id"], user["username"], req.session_token,
        req.question,
        result.get("answer", ""),
        result.get("sources", []),
        result.get("latency_ms", 0),
        result.get("confidence", "medium"),
    )
    result["log_id"] = log_id
    return result


# ── Feedback ───────────────────────────────────────────────────────────────
@app.post("/api/feedback")
async def feedback(req: FeedbackReq):
    user = get_current_user(req.session_token)
    if not user: raise HTTPException(401, "Invalid session")
    if req.feedback not in ("helpful", "not_helpful"):
        raise HTTPException(400, "feedback must be 'helpful' or 'not_helpful'")
    submit_feedback(req.log_id, req.feedback)
    return {"success": True}


# ── History ────────────────────────────────────────────────────────────────
@app.get("/api/history")
async def history(session_token: str):
    user = get_current_user(session_token)
    if not user: raise HTTPException(401, "Invalid session")
    return {"history": get_chat_history(session_token)}


# ── Audit Logs ─────────────────────────────────────────────────────────────
@app.get("/api/audit-logs")
async def audit_logs(session_token: str, limit: int = 500):
    user = get_current_user(session_token)
    if not user: raise HTTPException(401, "Invalid session")
    return {"logs": get_logs(limit)}


# ── Metrics ────────────────────────────────────────────────────────────────
@app.get("/api/metrics")
async def metrics(session_token: str):
    user = get_current_user(session_token)
    if not user: raise HTTPException(401, "Invalid session")
    return get_metrics()


# ── Document Ingest ────────────────────────────────────────────────────────
@app.post("/api/ingest")
async def ingest(req: IngestReq, session_token: str = ""):
    user = get_current_user(session_token)
    if not user or user["role"] != "admin":
        raise HTTPException(403, "Admin only")
    doc_id = rag.ingest(req.title, req.content, req.tags or "")
    return {"doc_id": doc_id, "message": "Document ingested successfully"}


# ── Documents ──────────────────────────────────────────────────────────────
@app.get("/api/documents")
async def documents():
    return {"documents": rag.list_docs()}


# ── User Management (Admin) ────────────────────────────────────────────────
@app.get("/api/users")
async def list_users(session_token: str):
    user = get_current_user(session_token)
    if not user or user["role"] != "admin": raise HTTPException(403, "Admin only")
    return {"users": get_all_users()}

@app.put("/api/users/{user_id}/role")
async def change_role(user_id: int, req: UserRoleReq):
    caller = get_current_user(req.session_token)
    if not caller or caller["role"] != "admin": raise HTTPException(403, "Admin only")
    update_user_role(user_id, req.role)
    return {"success": True}

@app.put("/api/users/{user_id}/active")
async def toggle_active(user_id: int, req: UserActiveReq):
    caller = get_current_user(req.session_token)
    if not caller or caller["role"] != "admin": raise HTTPException(403, "Admin only")
    set_user_active(user_id, req.is_active)
    return {"success": True}

@app.delete("/api/users/{user_id}")
async def remove_user(user_id: int, session_token: str):
    caller = get_current_user(session_token)
    if not caller or caller["role"] != "admin": raise HTTPException(403, "Admin only")
    delete_user(user_id)
    return {"success": True}


# ── Health ─────────────────────────────────────────────────────────────────
@app.get("/health")
async def health():
    return {
        "status":  "ok",
        "project": "DevOps Runbook Assistant",
        "model":   "llama-3.1-8b-instant via Groq",
        "vector":  "FAISS",
        "db_auth": "PostgreSQL (custom auth DB)",
        "db_app":  "PostgreSQL (custom app DB)",
        "chunks":  len(rag.store.metadata),
    }


if __name__ == "__main__":
    uvicorn.run("main:app", host="127.0.0.1", port=8000, reload=True)
